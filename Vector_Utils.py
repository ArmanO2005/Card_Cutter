from sentence_transformers import SentenceTransformer, util
import itertools
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.run import Run
import openai
import json
import spacy
from spacy.cli import download
from spacy.tokens import Span
import string
import re
import numpy as np


with open('secrets.json') as f:
    creds = json.load(f)


openai.api_key = creds["apikey"]
download("en_core_web_sm")
nlp = spacy.load("en_core_web_sm")



class vectorizer:
    def __init__(self, model_name="all-MiniLM-L6-v2"):
        self.model = SentenceTransformer(model_name)

    def encode(self, sentence):
        return self.model.encode(sentence, convert_to_tensor=True)
    
    def similarity(self, vector, prompt_vec):
        return util.pytorch_cos_sim(vector, prompt_vec).item()
    
    def screen_paragraph(self, sentences, prompt_vec, threshold=0.5):
        if not sentences:
            return False
        best_score = 0
        for i in sentences:
            vector = self.encode(i)
            sim_score = self.similarity(vector, prompt_vec)
            if sim_score > best_score:
                best_score = sim_score
        print(best_score)
        return best_score > threshold


#todo create subsets of a certain length or less
def ordered_combinations(iterable, length):
    items = list(iterable)
    return itertools.chain.from_iterable(
        itertools.combinations(items, r) for r in range(1, length+1)
    )


def greedy_combinations(iterable, length, vectorizer_instance, prompt_vec, threshold=0.3):
    items = list(iterable)
    pruned_items = [item for item in items if vectorizer_instance.similarity(vectorizer_instance.encode(item), prompt_vec) > threshold]
    return itertools.chain.from_iterable(
        itertools.combinations(pruned_items, r) for r in range(1, length+1)
    )



def read_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return "\n".join(full_text)


def llmd_prompt(prompt, article):
    full_prompt = "in 7-10 words, restate the prompt as best you can only using words from the article. Your response should should not include uncertain language. Your response should be no longer than 15 words: " + prompt + "\n\nArticle:\n" + article
    response = openai.ChatCompletion.create(
        model="gpt-4", 
        messages=[
            {"role": "system", "content": "You summarize and synthesize articles in a decisive and argumentative manner in 15 words or less."},
            {"role": "user", "content": full_prompt}
        ],
        max_tokens=150,
        temperature=0.5
    )
    return response.choices[0].message.content.strip()


#in 10-15 words, restate the following article contributions to making the argument in the prompt in the words of the article. Your response should should not include uncertain language. Your response should be no longer than 15 words: 
def paragraph_prompt(llmd_prompt, article):
    full_prompt = "restate the following article contributions to making the argument in the prompt in the words of the article. Your response should should not include uncertain language: " + llmd_prompt + "\n\nArticle:\n" + article
    response = openai.ChatCompletion.create(
        model="gpt-4", 
        messages=[
            {"role": "system", "content": "You summarize and synthesize articles in a decisive and argumentative manner"},
            {"role": "user", "content": full_prompt}
        ],
        max_tokens=150,
        temperature=0.5
    )
    return response.choices[0].message.content.strip()


def clause_separator(text):
    doc = nlp(text)
    clauses = []
    clause_tokens = []

    def flush_clause():
        if clause_tokens:
            span = Span(doc, clause_tokens[0].i, clause_tokens[-1].i + 1)
            text = span.text.strip()
            if len(text.split()) <= 3 and clauses:
                clauses[-1] += "[SPLIT]" + text
            else:
                clauses.append(text)
            clause_tokens.clear()


    inside_parentheses = False

    for token in doc:
        if token.text == "(":
            flush_clause()
            inside_parentheses = True
            continue
        elif token.text == ")":
            inside_parentheses = False
            flush_clause()
            continue

        clause_tokens.append(token)

        if inside_parentheses:
            continue

        if token.dep_ in ("ccomp", "advcl", "relcl", "conj") and token.head.dep_ != "ROOT":
            flush_clause()

        elif token.text in {".", ";", ":"}:
            flush_clause()
        elif token.text == ",":
            flush_clause()
        elif token.dep_ == "cc":  # coordinating conjunction like 'and', 'but'
            flush_clause()
        elif token.pos_ == "SCONJ" and token.text.lower() in {"although", "because", "if", "while"}:
            flush_clause()
        elif token.pos_ == "ADV" and token.text.lower() in {"however", "moreover", "therefore", "nevertheless"}:
            flush_clause()

    flush_clause()
    return clauses


def split_keep_delimiters(text, delimiters):
    escaped = [re.escape(d) for d in delimiters]
    pattern = f"({'|'.join(escaped)})"
    return re.split(pattern, text)


def underline_best_match_in_paragraph(paragraph, best_match, emphasis_clauses=None):
    if not best_match:
        small_text = paragraph.text
        paragraph.clear()
        paragraph.add_run(small_text).font.size = Pt(8)

    full_text = paragraph.text
    paragraph.clear()  # Clear existing text and runs

    current_index = 0

    for clause in best_match:
        clause_index = full_text.find(clause, current_index)

        if clause_index == -1:
            continue  # Skip if clause not found

        # Add the text before the match with size 8
        if clause_index > current_index:
            before_text = full_text[current_index:clause_index]
            run = paragraph.add_run(before_text)
            run.font.size = Pt(8)

        # Add the matched clause with underline and size 12
        if clause not in emphasis_clauses:
            run = paragraph.add_run(clause)
            run.underline = True
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(clause)
            run.underline = True
            run.font.bold = True
            run.font.size = Pt(11)
        #attempt at emphasis, works 
        

        current_index = clause_index + len(clause)

    # Add the remaining text after the last match
    if current_index < len(full_text):
        remaining = full_text[current_index:]
        run = paragraph.add_run(remaining)
        run.font.size = Pt(8)

def make_small(paragraph):
    full_text = paragraph.text
    paragraph.clear()
    run = paragraph.add_run(full_text)
    run.font.size = Pt(8)

def emphasis(paragraph, emphasis_words):
    if any(word in paragraph.text for word in emphasis_words):
        for run in paragraph.runs:
            if any(word in run.text for word in emphasis_words):
                run.font.bold = True
                run.underline = True
                run.font.size = Pt(11)
